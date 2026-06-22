from __future__ import annotations

import os
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..config import get_settings
from ..models import Contract, ContractVersion, Language

# A Unicode font that covers Devanagari so Marathi prints correctly.
DEVANAGARI_FONT = "Nirmala UI"
LATIN_FONT = "Calibri"


def _set_run_fonts(run, latin: str, complex_script: str) -> None:
    """python-docx doesn't expose the complex-script (Devanagari) font; set it via XML."""
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), complex_script)  # complex script = Devanagari


def _para(doc, text: str, *, size: int = 11, bold: bool = False,
          italic: bool = False, color: Optional[RGBColor] = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _set_run_fonts(run, LATIN_FONT, DEVANAGARI_FONT)
    return p


def render_docx(contract: Contract, version: ContractVersion,
                language: Language = Language.en) -> str:
    """Render one contract version to a DOCX file and return its path."""
    settings = get_settings()
    storage = settings.contract_mint_storage_dir
    os.makedirs(storage, exist_ok=True)

    lang = language.value
    doc = Document()

    title = _para(doc, contract.title.upper(), size=16, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para(doc, f"Version {version.version}    Jurisdiction: India    Language: {lang.upper()}",
          size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66)).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    for i, clause in enumerate(version.clauses, start=1):
        _para(doc, f"{i}. {clause.title}", size=12, bold=True)
        body = clause.body.get(lang) or clause.body.get("en", "")
        _para(doc, body, size=11)
        plain = clause.plain.get(lang) or clause.plain.get("en", "")
        if plain:
            _para(doc, f"In plain words: {plain}", size=9, italic=True,
                  color=RGBColor(0x33, 0x66, 0x99))
        doc.add_paragraph()

    doc.add_paragraph()
    _para(doc,
          "This document was prepared with an assistive tool and is not legal advice. "
          "Please review before signing. The merchant confirms they have reviewed this version.",
          size=8, italic=True, color=RGBColor(0x99, 0x33, 0x33))

    filename = f"{contract.contract_id}_v{version.version}_{lang}.docx"
    path = os.path.join(storage, filename)
    doc.save(path)
    return path
